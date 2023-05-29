import urllib.request
from bs4 import BeautifulSoup
import docx
import time
import re

header = {
    "Referer": "http://www.gov.cn/",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/111.0.0.0 Safari/537.36"
}

def getSoup(url):
    page = urllib.request.Request(url, headers=header)
    page = urllib.request.urlopen(page).read()
    return BeautifulSoup(page)

def getNumPages(url):
    soup = getSoup(url)
    numPages = re.findall("\d+", soup.find("div", {'class': "newspage cl"}).find_all('li')[11].text)[0]
    numPages = int(numPages)
    return numPages

def getArticleList(url):
    numPages = getNumPages(url)
    articleList = []
    subArticles = []

    for page in range(numPages):
        url = "http://sousuo.gov.cn/column/30562/" + str(page) + ".htm"
        soup = getSoup(url)
        articles = soup.find('ul', {'class': 'listTxt'}).findAll('h4')
        articles[0].find('span').text

        for article in articles:
            articleList.append(article.find('a').get('href'))
    
        time.sleep(1)
    
    for article in articleList:
        soup = getSoup(article)
        for link in soup.find_all('a'):
            if "http://www.gov.cn/premier/" in str(link.get('href')):
                subArticles.append(link.get('href'))

        #time.sleep(1)
    
    return [*set(subArticles)]

def writeArticle(url):
    doc = docx.Document()

    page = urllib.request.Request(url)
    page = urllib.request.urlopen(page).read()
    soup = BeautifulSoup(page)

    title = soup.find("div", {"class": "article oneColumn pub_border"})
    title = title.find("h1").text.strip()
    date = soup.find("div", {"class": "pages-date"}).text
    date = date.split("来源")[0]

    doc.add_heading(date + ' - ' + title)

    for txt in soup.find_all('p', {'style':'text-indent: 2em; font-family: 宋体; font-size: 12pt;'}):
        doc.add_paragraph(txt.text)
        
    doc.save("china_crawling/" + date + ' - ' + title + ".docx")

if __name__ == "__main__":
    base_url = "http://sousuo.gov.cn/column/30562/0.htm"
    article_urls = getArticleList(base_url)
    for article in article_urls:
        writeArticle(article)